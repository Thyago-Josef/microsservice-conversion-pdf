import os
import uuid
import subprocess
import re
import fitz
from flask import Flask, request, jsonify, send_file
from docling.document_converter import DocumentConverter
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz
from PIL import Image
import io
from lxml import etree






app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
PANDOC_PATH = r"C:\pandoc-3.9.0.2\pandoc.exe"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

converter = DocumentConverter()



def crop_icon_to_content(img_bytes):
    """
    Remove o fundo branco ao redor do ícone (torna transparente)
    e corta só a área com conteúdo real.
    """
    img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    
    # Converte fundo branco em transparente
    data = img.getdata()
    new_data = []
    for r, g, b, a in data:
        if r > 220 and g > 220 and b > 220:
            new_data.append((255, 255, 255, 0))
        else:
            new_data.append((r, g, b, a))
    img.putdata(new_data)
    
    # Corta a área com conteúdo ignorando pixels transparentes
    bbox = img.getbbox()
    if bbox:
        img = img.crop(bbox)
    
    # Adiciona pequeno padding
    padded = Image.new("RGBA", (img.width + 4, img.height + 4), (255, 255, 255, 0))
    padded.paste(img, (2, 2))
    
    out = io.BytesIO()
    padded.save(out, format="PNG")
    return out.getvalue()

def clean_oriental_chars(docx_path):
    """Remove apenas caracteres chineses/japoneses (mantém ícones em Hangul)"""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue
            cleaned = []
            for c in text:
                cp = ord(c)
                # Remove CJK Unified Ideographs e Japanese Kana, mantém Hangul (ícones)
                if not (0x4E00 <= cp <= 0x9FFF or 0x3040 <= cp <= 0x30FF):
                    cleaned.append(c)
            new_text = ''.join(cleaned)
            if new_text != text:
                para.text = new_text
        doc.save(docx_path)
    except Exception as e:
        print(f"⚠️ Erro ao limpar caracteres: {e}")


def extract_icon_images_from_pdf(pdf_path):
    """Extrai ícones do PDF renderizando as áreas onde aparecem caracteres Hangul"""
    doc = fitz.open(pdf_path)
    icon_map = {}  # char -> img_bytes
    hangul_spans_found = 0
    
    # Cria diretório para debug
    debug_dir = os.path.join(OUTPUT_FOLDER, "debug_icons")
    os.makedirs(debug_dir, exist_ok=True)
    
    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"]
                    if not text:
                        continue
                    
                    has_hangul = any(0xAC00 <= ord(c) <= 0xD7AF for c in text)
                    
                    if has_hangul:
                        hangul_spans_found += 1
                        bbox = span["bbox"]
                        
                        # Expande a bbox
                        x0 = max(0, bbox[0] - 2)
                        y0 = max(0, bbox[1] - 2)
                        x1 = min(page.rect.width, bbox[2] + 2)
                        y1 = min(page.rect.height, bbox[3] + 2)
                        
                        rect = fitz.Rect(x0, y0, x1, y1)
                        zoom = 6
                        mat = fitz.Matrix(zoom, zoom)
                        pix = page.get_pixmap(matrix=mat, clip=rect)
                        
                        img_bytes = pix.tobytes("png")
                        
                        # Salva para debug
                        debug_path = os.path.join(debug_dir, f"icon_{ord(text[0]):04X}.png")
                        pix.save(debug_path)
                        print(f"Ícone salvo: {debug_path}")
                        
                        # Mapeia cada caractere Hangul para a imagem
                        for c in text:
                            if 0xAC00 <= ord(c) <= 0xD7AF and c not in icon_map:
                                icon_map[c] = img_bytes
                                print(f"Ícone mapeado: '{c}' (U+{ord(c):04X})")
    
    doc.close()
    print(f"Total de spans com Hangul: {hangul_spans_found}")
    print(f"Total de ícones únicos: {len(icon_map)}")
    return icon_map


# def replace_text_with_icons_in_docx(docx_path, icon_map):
#     """Substitui cada caractere Hangul no DOCX pela imagem do ícone correspondente"""
#     try:
#         doc = Document(docx_path)
        
#         if not icon_map:
#             print("ℹ️ Nenhum ícone para substituir")
#             return
        
#         print(f"🗺️ Ícones disponíveis: {[f'U+{ord(k):04X}' for k in icon_map.keys()]}")
        
#         # Processa parágrafos
#         for para in doc.paragraphs:
#             para_text = para.text
            
#             if not para_text:
#                 continue
            
#             # Verifica se há Hangul para substituir
#             has_hangul = any(0xAC00 <= ord(c) <= 0xD7AF and c in icon_map for c in para_text)
#             if not has_hangul:
#                 continue
            
#             print(f"🔄 Processando: '{para_text[:60]}...'")
            
#             # Salva a formatação do parágrafo (alinhamento, espaçamento)
#             align = para.paragraph_format.alignment
#             space_before = para.paragraph_format.space_before
#             space_after = para.paragraph_format.space_after
            
#             # Reconstrói o parágrafo
#             # Limpa todos os runs
#             for run in para.runs:
#                 run._r.getparent().remove(run._r)
            
#             # Adiciona novo conteúdo
#             current_text = ""
#             for c in para_text:
#                 if 0xAC00 <= ord(c) <= 0xD7AF and c in icon_map:
#                     # Adiciona texto acumulado
#                     if current_text:
#                         para.add_run(current_text)
#                         current_text = ""
                    
#                     # Insere ícone
#                     try:
#                         img_bytes = icon_map[c]
#                         img_stream = io.BytesIO(img_bytes)
#                         run = para.add_run()
#                         run.add_picture(img_stream, width=Inches(0.18))
#                         print(f"   ✅ Ícone U+{ord(c):04X} inserido")
#                     except Exception as e:
#                         print(f"   ❌ Erro: {e}")
#                         para.add_run(c)
#                 else:
#                     current_text += c
            
#             # Adiciona texto restante
#             if current_text:
#                 para.add_run(current_text)
            
#             # Restaura formatação
#             para.paragraph_format.alignment = align
#             if space_before:
#                 para.paragraph_format.space_before = space_before
#             if space_after:
#                 para.paragraph_format.space_after = space_after
        
#         doc.save(docx_path)
#         print(f"✅ Ícones substituídos no DOCX")
#     except Exception as e:
#         print(f"Erro ao processar DOCX: {e}")
#         import traceback
#         traceback.print_exc()

def split_merged_paragraphs(docx_path):
    """
    Quebra parágrafos que contêm múltiplos bullets num único <w:p>
    Trata dois casos:
    1. Bullets separados por <w:br/> 
    2. Bullets concatenados no mesmo <w:r> ou em <w:r> consecutivos
    """
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    for para in list(doc.paragraphs):
        p_elem = para._p
        runs = p_elem.findall(f'{{{W}}}r')

        if not runs:
            continue

        # Verifica se tem <w:br/> ou múltiplos bullets (•) nos runs
        has_br = any(r.find(f'{{{W}}}br') is not None for r in runs)
        full_text = para.text
        bullet_count = full_text.count('•')

        if not has_br and bullet_count <= 1:
            continue

        print(f"  Quebrando: '{full_text[:60]}...'")

        parent = p_elem.getparent()
        insert_pos = list(parent).index(p_elem)
        original_pPr = p_elem.find(f'{{{W}}}pPr')

        # Agrupa runs em segmentos separados por <w:br/> ou por início de bullet
        segments = []
        current_runs = []

        for elem in p_elem:
            tag = elem.tag.split('}')[1] if '}' in elem.tag else elem.tag

            if tag == 'pPr':
                continue

            if tag == 'r':
                br = elem.find(f'{{{W}}}br')
                t_elem = elem.find(f'{{{W}}}t')
                t_text = t_elem.text if t_elem is not None else ''

                if br is not None:
                    # <w:br/> → fecha segmento atual
                    if current_runs:
                        segments.append(current_runs)
                    current_runs = []
                elif t_text and '•' in t_text:
                    # Novo bullet → pode haver múltiplos bullets no mesmo run
                    # Separa cada bullet em seu próprio segmento
                    parts = [p for p in t_text.split('•') if p.strip()]
                    for i, part in enumerate(parts):
                        if current_runs and i == 0:
                            # Fecha segmento anterior
                            segments.append(current_runs)
                            current_runs = []
                        # Cria run com este bullet
                        new_r = deepcopy(elem)
                        new_t = new_r.find(f'{{{W}}}t')
                        if new_t is not None:
                            new_t.text = '• ' + part.strip()
                        current_runs = [new_r]
                        if i < len(parts) - 1:
                            segments.append(current_runs)
                            current_runs = []
                else:
                    current_runs.append(deepcopy(elem))
            else:
                current_runs.append(deepcopy(elem))

        if current_runs:
            segments.append(current_runs)

        if len(segments) <= 1:
            continue

        # Cria um parágrafo novo para cada segmento
        new_paras = []
        for seg_runs in segments:
            if not any(
                (r.find(f'{{{W}}}t') is not None and r.find(f'{{{W}}}t').text or '').strip()
                for r in seg_runs
            ):
                continue

            new_p = OxmlElement('w:p')
            if original_pPr is not None:
                new_p.append(deepcopy(original_pPr))
            for r in seg_runs:
                new_p.append(r)
            new_paras.append(new_p)

        if not new_paras:
            continue

        # Substitui parágrafo original pelos novos
        parent.remove(p_elem)
        for i, new_p in enumerate(new_paras):
            parent.insert(insert_pos + i, new_p)
        print(f"    → {len(new_paras)} parágrafos criados")

    doc.save(docx_path)
    print("✅ Parágrafos quebrados com sucesso")

def remove_picture_border(run):
    """Remove borda que LibreOffice/Word adiciona em imagens inline"""
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }
    for spPr in run._r.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}spPr'):
        # Remove bordas existentes
        for ln in spPr.findall('{http://schemas.openxmlformats.org/drawingml/2006/main}ln'):
            spPr.remove(ln)
        # Adiciona ln com noFill (sem borda)
        ln = etree.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}ln')
        etree.SubElement(ln, '{http://schemas.openxmlformats.org/drawingml/2006/main}noFill')
        break



def replace_text_with_icons_in_docx(docx_path, icon_map):
    try:
        doc = Document(docx_path)
        
        if not icon_map:
            print("ℹ️ Nenhum ícone para substituir")
            return
        
        for para in doc.paragraphs:
            para_text = para.text
            if not para_text:
                continue
            
            has_hangul = any(0xAC00 <= ord(c) <= 0xD7AF and c in icon_map for c in para_text)
            if not has_hangul:
                continue

            align = para.paragraph_format.alignment
            space_before = para.paragraph_format.space_before
            space_after = para.paragraph_format.space_after
            
            for run in para.runs:
                run._r.getparent().remove(run._r)
            
            current_text = ""
            for c in para_text:
                if 0xAC00 <= ord(c) <= 0xD7AF and c in icon_map:
                    if current_text:
                        para.add_run(current_text)
                        current_text = ""
                    
                    try:
                        img_bytes = icon_map[c]
                        
                        # ✅ APLICA crop_icon_to_content AQUI
                        clean_img = crop_icon_to_content(img_bytes)
                        
                        img_stream = io.BytesIO(clean_img)  # ← era icon_map[c]
                        run = para.add_run()
                        run.add_picture(img_stream, width=Inches(0.11))
                        remove_picture_border(run)
                    except Exception as e:
                        print(f"   ❌ Erro: {e}")
                        para.add_run(c)
                else:
                    current_text += c
            
            if current_text:
                para.add_run(current_text)
            
            para.paragraph_format.alignment = align
            if space_before:
                para.paragraph_format.space_before = space_before
            if space_after:
                para.paragraph_format.space_after = space_after
        
        doc.save(docx_path)
    except Exception as e:
        print(f"Erro ao processar DOCX: {e}")
        import traceback
        traceback.print_exc()




# -----------------------------------------------------------------------------------------

def extract_lines_from_pdf(pdf_path):
    """Extrai coordenadas das linhas/molduras do PDF"""
    doc = fitz.open(pdf_path)
    lines = []
    
    for page_num, page in enumerate(doc):
        # Extrai shapes (linhas, retângulos)
        shapes = page.get_drawings()
        for shape in shapes:
            rect = shape.get("rect", None)
            if rect and (abs(rect.width) > 20 or abs(rect.height) > 20):
                lines.append({
                    "page": page_num,
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1,
                    "width": rect.width,
                    "height": rect.height,
                    "color": shape.get("color", (0,0,0)),
                    "fill": shape.get("fill", None)
                })
        
        # Extrai linhas via annots
        annot = page.first_annot
        while annot:
            if annot.type == 2:  # Line
                line = annot.line
                if line:
                    lines.append({
                        "page": page_num,
                        "x0": line[0].x,
                        "y0": line[0].y,
                        "x1": line[1].x,
                        "y1": line[1].y
                    })
            annot = annot.next
    
    doc.close()
    return lines


def add_line_to_docx(doc, x0, y0, x1, y1, page_height, page_width):
    """Adiciona uma linha no DOCX usando tabela sem borda"""
    try:
        width_pt = abs(x1 - x0)
        height_pt = abs(y1 - y0)
        
        if width_pt < 20 and height_pt < 20:
            return
        
        # Cria tabela 1x1 para linha
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.cell(0, 0)
        
        # Determina espessura baseada no tamanho
        if width_pt > height_pt:  # Linha horizontal
            cell.text = "─" * int(min(width_pt / 10, 50))
        else:  # Linha vertical
            cell.text = "│"
        
        # Remove bordas
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find_or_add('w:tcBorders')
        
        # Define borda inferior para linha horizontal
        if width_pt > height_pt:
            for border in ['top', 'bottom']:
                element = OxmlElement(f'w:{border}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')
                tcBorders.append(element)
        
    except Exception as e:
        pass  # Silencia erros


def add_section_borders(docx_path, pdf_path):
    from collections import defaultdict
    from docx.oxml.ns import qn
    from docx.shared import Pt

    # 1. Extrai linhas gráficas reais do PDF
    doc_pdf = fitz.open(pdf_path)
    pdf_lines = []  # lista de {"page", "y", "width"}

    for page_num, page in enumerate(doc_pdf):
        page_height = page.rect.height
        for shape in page.get_drawings():
            rect = shape.get("rect", None)
            if rect and abs(rect.width) > 100 and abs(rect.height) < 3:
                pdf_lines.append({
                    "page": page_num,
                    "y": round(rect.y0, 1),
                    "page_height": page_height
                })

    # Agrupa linhas duplicadas (linhas duplas próximas viram uma)
    grouped = []
    for line in sorted(pdf_lines, key=lambda l: (l["page"], l["y"])):
        if grouped and grouped[-1]["page"] == line["page"] and abs(line["y"] - grouped[-1]["y"]) < 5:
            # média do Y do par de linhas duplas
            grouped[-1]["y"] = round((grouped[-1]["y"] + line["y"]) / 2, 1)
        else:
            grouped.append(dict(line))

    if not grouped:
        print("📏 PDF sem linhas horizontais, pulando bordas")
        doc_pdf.close()
        return

    # Detecta se são linhas duplas
    raw_ys = [round(l["y"]) for l in pdf_lines]
    is_double = any(
        abs(sorted(set(raw_ys))[i+1] - sorted(set(raw_ys))[i]) < 5
        for i in range(len(sorted(set(raw_ys))) - 1)
    )
    border_style = "double" if is_double else "single"
    print(f"📏 Estilo: {border_style} | {len(grouped)} linhas encontradas")

    # 2. Para cada linha, acha o parágrafo acima e abaixo no PDF
    #    e calcula a proporção (onde a linha está entre os dois)
    border_targets = []

    for line in grouped:
        page = doc_pdf[line["page"]]
        all_spans = []
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for pdf_line in block["lines"]:
                for span in pdf_line["spans"]:
                    text = span["text"].strip()
                    if text:
                        all_spans.append({
                            "text": text,
                            "y_top": round(span["bbox"][1], 1),    # topo do texto
                            "y_bot": round(span["bbox"][3], 1),    # base do texto
                            "size": round(span["size"], 1),
                        })

        line_y = line["y"]

        # Parágrafo cujo FUNDO (y_bot) está acima da linha
        above = [s for s in all_spans if s["y_bot"] < line_y]
        # Parágrafo cujo TOPO (y_top) está abaixo da linha
        below = [s for s in all_spans if s["y_top"] > line_y]

        if not above or not below:
            continue

        para_above = max(above, key=lambda s: s["y_bot"])
        para_below = min(below, key=lambda s: s["y_top"])

        # Distância real entre o fundo do parágrafo acima e o topo do parágrafo abaixo
        gap_total   = para_below["y_top"] - para_above["y_bot"]  # espaço total entre os dois
        dist_above  = line_y - para_above["y_bot"]               # linha → fundo do parágrafo acima
        dist_below  = para_below["y_top"] - line_y               # linha → topo do parágrafo abaixo

        # Proporção: 0.0 = colada acima, 1.0 = colada abaixo, 0.5 = centralizada
        ratio = dist_above / gap_total if gap_total > 0 else 0.5

        # Decide onde ancorar baseado na fonte (título tem size maior)
        if para_below["size"] >= para_above["size"]:
            anchor = "top"           # borda superior do parágrafo abaixo
            anchor_text = para_below["text"].upper()[:40]
            # space = distância do topo do título até a linha (em pontos)
            space_pt = round(dist_below)
        else:
            anchor = "bottom"        # borda inferior do parágrafo acima
            anchor_text = para_above["text"].upper()[:40]
            space_pt = round(dist_above)

        # Word usa unidades de 1/8 pt para w:space em bordas de parágrafo
        # Máximo permitido é 31 (= ~4pt). Normaliza para esse range.
        space_val = min(31, max(1, round(space_pt / 2)))

        border_targets.append({
            "anchor": anchor,
            "text": anchor_text,
            "space": str(space_val),
            "ratio": round(ratio, 2),
            "dist_above_pt": round(dist_above, 1),
            "dist_below_pt": round(dist_below, 1),
        })
        print(f"📏 Y={line_y} | gap={round(gap_total)}pt | acima={round(dist_above)}pt abaixo={round(dist_below)}pt ratio={round(ratio,2)} → {anchor} em '{anchor_text[:25]}' space={space_val}")

    doc_pdf.close()

    # 3. Aplica no DOCX
    doc = Document(docx_path)
    applied = set()

    for para in doc.paragraphs:
        text = para.text.strip().upper()
        if not text:
            continue

        for target in border_targets:
            match_text = target["text"][:20]
            if not match_text or match_text in applied:
                continue
            if not text.startswith(match_text):
                continue

            p = para._p
            pPr = p.get_or_add_pPr()

            # Remove borda existente
            existing = pPr.find(qn('w:pBorders'))
            if existing is not None:
                pPr.remove(existing)

            # Adiciona borda com espaçamento proporcional ao PDF
            pBorders = OxmlElement('w:pBorders')
            border_el = OxmlElement(
                'w:bottom' if target["anchor"] == "bottom" else 'w:top'
            )
            border_el.set(qn('w:val'), border_style)
            border_el.set(qn('w:sz'), '6')
            border_el.set(qn('w:space'), target["space"])
            border_el.set(qn('w:color'), '000000')
            pBorders.append(border_el)
            pPr.append(pBorders)

            # Ajusta space_before/after do parágrafo para replicar o gap do PDF
            if target["anchor"] == "top":
                para.paragraph_format.space_before = Pt(target["dist_above_pt"])
            else:
                para.paragraph_format.space_after = Pt(target["dist_below_pt"])

            applied.add(match_text)
            print(f"   ✅ '{target['anchor']}' em '{text[:30]}' | space_before/after={target['dist_above_pt' if target['anchor'] == 'top' else 'dist_below_pt']}pt")
            break

    doc.save(docx_path)
    print(f"✅ {len(applied)} bordas adicionadas")

def fix_title_size(docx_path, doc_info):
    """Usa o tamanho de título extraído do PDF pelo analyzer"""
    title_sizes = doc_info.get("title_sizes", [])
    
    if not title_sizes:
        return  # sem info, não mexe
    
    # title_sizes já vem ordenado do maior para menor
    # o primeiro é o título principal (nome)
    title_pt = title_sizes[0]
    
    doc = Document(docx_path)
    for para in doc.paragraphs:
        if para.text.strip():
            for run in para.runs:
                run.font.size = Pt(title_pt)
            break
    doc.save(docx_path)



def convert_pdf_pdf2docx(pdf_path, docx_path):
    """Converte PDF para DOCX mantendo formatação"""
    print(f"🎯 Convertendo via pdf2docx: {pdf_path}")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

    # ✅ Primeira coisa: quebra parágrafos agrupados
    print("✂️ Quebrando parágrafos agrupados...")
    split_merged_paragraphs(docx_path)


    # ✅ Analisa o PDF para pegar tamanhos reais
    from analyzer import analyze_pdf, extract_doc_info
    styles, icons = analyze_pdf(pdf_path)
    doc_info = extract_doc_info(styles)

    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

    # Extrai ícones do PDF
    print("🎨 Extraindo ícones do PDF...")
    icon_map = extract_icon_images_from_pdf(pdf_path)
    
    # Substitui ícones no DOCX ANTES de limpar caracteres
    if icon_map:
        print(f"📸 {len(icon_map)} ícones extraídos, substituindo no DOCX...")
        replace_text_with_icons_in_docx(docx_path, icon_map)
    else:
        print("ℹ️ Nenhum ícone encontrado no PDF")
    
    # Limpa caracteres orientais DEPOIS de substituir ícones
    print("🧹 Limpando caracteres orientais...")
    clean_oriental_chars(docx_path)
    
    # Adiciona bordas nas seções (linhas horizontais)
    print("📏 Adicionando bordas nas seções...")
    add_section_borders(docx_path, pdf_path)
    print("🔤 Ajustando tamanho do título...")
    fix_title_size(docx_path, doc_info)

    print(f"✅ DOCX gerado: {docx_path}")


def clean_markdown(md_path):
    """Limpa caracteres orientais do Markdown"""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    cleaned = []
    for char in content:
        cp = ord(char)
        if (
            cp <= 0x024F
            or 0x0300 <= cp <= 0x036F
            or char in '\n\r\t'
            or 0x1F000 <= cp <= 0x1FFFF
            or 0x2600 <= cp <= 0x27BF
        ):
            cleaned.append(char)

    content = "".join(cleaned)

    lines = content.split('\n')
    filtered_lines = []
    for line in lines:
        stripped = line.strip()
        has_real_text = any(c.isalpha() for c in stripped)

        if stripped == '':
            filtered_lines.append('')
        elif stripped.startswith('#'):
            filtered_lines.append(line)
        elif stripped == '---':
            filtered_lines.append(line)
        elif has_real_text:
            filtered_lines.append(line)

    content = '\n'.join(filtered_lines)
    content = re.sub(r'\n{3,}', '\n\n', content)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content.strip())

    print("✅ Markdown limpo")



    

# ==================== ROTAS ====================

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "UP", "service": "docling-converter"})


@app.route("/convert/docx", methods=["POST"])
def convert_to_docx():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    strategy = request.form.get("strategy", "pdf2docx")

    file = request.files["file"]
    file_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_FOLDER, f"{file_id}_{file.filename}")
    file.save(input_path)

    try:
        docx_filename = f"{file_id}.docx"
        docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)

        if strategy == "pdf2docx":
            convert_pdf_pdf2docx(input_path, docx_path)

        elif strategy == "docling":
            from analyzer import analyze_pdf, extract_doc_info
            from reference_builder import build_reference_docx

            print(f"🎯 Convertendo via Docling + Pandoc: {input_path}")

            styles, icons = analyze_pdf(input_path)
            doc_info = extract_doc_info(styles)

            result = converter.convert(input_path)
            markdown_content = result.document.export_to_markdown()

            md_filename = f"{file_id}.md"
            md_path = os.path.join(OUTPUT_FOLDER, md_filename)
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)

            clean_markdown(md_path)

            ref_path = build_reference_docx(doc_info, OUTPUT_FOLDER)
            lang = doc_info.get("language", "pt-BR")

            result_pandoc = subprocess.run([
                PANDOC_PATH,
                md_path,
                "-o", docx_path,
                "--from", "markdown",
                "--to", "docx",
                "--standalone",
                f"--reference-doc={ref_path}",
                "--metadata", f"lang={lang}"
            ], capture_output=True, text=True)

            if result_pandoc.returncode != 0:
                raise RuntimeError(f"Pandoc falhou: {result_pandoc.stderr}")

            os.remove(md_path)
            os.remove(ref_path)

            print(f"✅ DOCX gerado: {docx_path}")

        else:
            return jsonify({"error": "Estratégia inválida. Use: pdf2docx ou docling"}), 400

        return jsonify({
            "status": "success",
            "file_id": file_id,
            "strategy": strategy,
            "download_url": f"/download/{docx_filename}"
        })

    except Exception as e:
        print(f"❌ Erro: {str(e)}")
        return jsonify({"error": str(e)}), 500

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)


@app.route("/convert/markdown", methods=["POST"])
def convert_to_markdown():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    from analyzer import analyze_pdf, extract_doc_info

    file = request.files["file"]
    file_id = str(uuid.uuid4())
    input_path = os.path.join(UPLOAD_FOLDER, f"{file_id}_{file.filename}")
    file.save(input_path)

    try:
        styles, icons = analyze_pdf(input_path)
        doc_info = extract_doc_info(styles)

        result = converter.convert(input_path)
        markdown_content = result.document.export_to_markdown()

        md_filename = f"{file_id}.md"
        md_path = os.path.join(OUTPUT_FOLDER, md_filename)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)

        clean_markdown(md_path)

        return jsonify({
            "status": "success",
            "file_id": file_id,
            "download_url": f"/download/{md_filename}",
            "doc_info": doc_info
        })

    except Exception as e:
        print(f"❌ Erro: {str(e)}")
        return jsonify({"error": str(e)}), 500

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)


@app.route("/download/<path:filename>", methods=["GET"])
def download(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        return jsonify({"error": "Arquivo não encontrado"}), 404
    return send_file(path, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)